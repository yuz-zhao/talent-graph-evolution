"""
生成《TalentGraph Evolution 知识图谱设计文档 V4.0》Word 版本
基于赛题 XH-202621 要求 + V3 设计方案 + 实际数据资产
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

OUTPUT = os.path.join(os.path.dirname(__file__), '..', '..', '..', '文档',
                       'TalentGraph_知识图谱设计文档_V4_最终版.docx')

doc = Document()

# ============================================================
# 样式设置
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_font = heading_style.font
    heading_font.name = '微软雅黑'
    heading_font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_para(text, bold=False, size=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    # data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()
    return table

# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

add_para('TalentGraph Evolution', bold=True, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('知识图谱设计文档 V4.0', bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para('多源异构数据驱动的岗位能力图谱构建与动态演化', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('—— 基于 Neo4j 图数据库的中文知识图谱方案 ——', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()
add_para('赛题编号：XH-202621', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('文档日期：2026年7月', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('版本状态：最终版（可直接实施）', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ============================================================
# 目录页
# ============================================================
doc.add_heading('目录', level=1)
toc_items = [
    '一、系统定位与闭环设计',
    '二、数据资产概述',
    '三、知识图谱总体架构',
    '四、核心节点设计（12类）',
    '五、核心关系设计（18种）',
    '六、Neo4j 建图规范与 Cypher 示例',
    '七、数据到图谱的映射流程',
    '八、新岗位发现算法设计',
    '九、既有岗位能力动态演化',
    '十、人岗匹配与能力差距分析',
    '十一、GraphRAG 证据链与防幻觉设计',
    '十二、学习路径推荐',
    '十三、评估指标体系',
    '十四、分阶段实施路线图',
    '附录A：节点属性完整清单',
    '附录B：Cypher 查询参考',
    '附录C：数据文件清单',
]
for item in toc_items:
    add_para(item, size=11)

doc.add_page_break()

# ============================================================
# 第一章：系统定位与闭环设计
# ============================================================
doc.add_heading('一、系统定位与闭环设计', level=1)

doc.add_heading('1.1 赛题核心要求', level=2)
add_para('赛题 XH-202621 要求构建一套基于人工智能的岗位能力动态分析系统，核心考察以下能力：')
add_para('① 新岗位发现与定义：识别市场中涌现、萌芽、尚未标准化的新兴岗位，生成岗位定义（名称、职责、必备技能、加分技能、典型行业应用场景），支持人工优化与动态更新。')
add_para('② 既有岗位能力动态更新：识别既有岗位（如Java开发工程师）能力要求的变化，提供变化说明和来源证据，标注技能新增/删除/修改等操作，支持人工优化与动态演化。')
add_para('③ 新一代信息技术岗位全息图谱：展示所有岗位及其能力要求，可视化技能热度、技能图谱，支持人岗匹配诊断、简历解析（PDF/Word）、多维匹配分析和学习路径规划。')
add_para('④ 质量指标：JD解析准确率≥90%，简历提取准确率≥90%，匹配准确率≥90%，测试集≥100条。')

doc.add_heading('1.2 全流程闭环', level=2)
add_para('赛题评分标准明确指出"全流程闭环系统"为最高评分项（30分）。本系统闭环设计如下：', bold=True)

add_table(
    ['阶段', '输入', '处理', '输出', '验证'],
    [
        ['① 数据采集', '招聘网站/企业官网/GitHub/arXiv/ Blog', '多源异构数据清洗、去重、标准化', '结构化岗位JD、简历、技术趋势数据', 'source_type标注、证据追溯'],
        ['② 新岗位发现', '清洗后JD + 技术趋势数据', '技能共现聚类 → 岗位簇识别 → LLM生成定义 → 证据校验 → 人工确认', '新岗位候选节点 + 岗位定义', '稳定性评分、多源交叉验证'],
        ['③ 既有岗位演化', '多期JD数据 + 技能本体', '时间切片对比 → 变化检测 → LLM解释 → 证据标注', '能力变更事件节点', 'support_count、source_ diversity'],
        ['④ 知识图谱构建', '全部结构化数据', 'ETL映射 → Neo4j导入 → 关系建立 → 证据关联', '岗位能力全息图谱', '节点唯一性、关系完整性'],
        ['⑤ 简历解析', 'PDF/Word简历', '文档解析 → 要素提取 → 技能标准化 → 画像生成', '结构化人才画像', '提取准确率≥90%'],
        ['⑥ 人岗匹配', '人才画像 + 岗位图谱', '多维匹配 → 差距分析 → 推荐解释', '匹配报告 + 能力差距', '匹配准确率≥90%'],
        ['⑦ 学习路径', '能力差距 + 课程/证书库', '差距排序 → 资源召回 → 路径规划', '阶梯式学习计划', '课程-技能映射覆盖率'],
        ['⑧ GraphRAG证据', '用户问题 + 图谱 + 向量库', '图路径检索 + 文本召回 → LLM整合 → 证据标注', '可解释的回答', '证据命中率≥90%'],
    ]
)

add_para('闭环示意：', bold=True)
add_para('数据采集 → 新岗位发现 → 图谱构建 → 简历解析 → 人岗匹配 → 能力差距 → 学习路径 → 能力提升 → 重新匹配（循环）')

doc.add_page_break()

# ============================================================
# 第二章：数据资产概述
# ============================================================
doc.add_heading('二、数据资产概述', level=1)

doc.add_heading('2.1 数据文件清单', level=2)
add_table(
    ['层级', '文件名', '规模', '来源类型', '用途'],
    [
        ['核心', 'jd_clean.csv', '2,500条', 'real_crawled + processed', '岗位节点、公司节点、岗位-技能关系'],
        ['核心', 'resume_clean.csv', '500条', 'github_public + synthetic', '人才节点、项目节点、人才-技能关系'],
        ['核心', 'skill_ontology.json', '545个技能', 'processed', '技能节点、技能分类、技能层级'],
        ['核心', 'job_standard_dict.csv', '37条', 'processed', '岗位标准化映射'],
        ['证据', 'github_detail.jsonl', '389条', 'public_api', '技术项目节点、技术趋势证据'],
        ['证据', 'github_trend.jsonl', '389条', 'public_api', '技术热度数据'],
        ['证据', 'arxiv_trend.jsonl', '359条', 'public_api', '学术论文节点'],
        ['证据', 'blog_trend.jsonl', '857条', 'public_api/RSS', '技术文章节点'],
        ['教育', 'course_data.csv', '384条', 'template_generated', '课程节点、学习资源'],
        ['教育', 'certificate_data.csv', '101条', 'template_generated', '证书节点'],
        ['评估', 'gold_jd_set_reviewed.json', '100条', 'manual_reviewed', 'JD解析准确率评估'],
        ['评估', 'gold_resume_set_reviewed.json', '50条', 'manual_reviewed', '简历提取准确率评估'],
        ['评估', 'match_label_set_reviewed.json', '100条', 'manual_reviewed', '匹配准确率评估'],
        ['评估', 'negative_samples_reviewed.json', '50条', 'manual_reviewed', '异常检测评估'],
    ]
)

doc.add_heading('2.2 数据质量概览', level=2)
add_table(
    ['数据文件', '关键字段', '填充率', '说明'],
    [
        ['jd_clean.csv', 'industry / publish_time / skill_standard', '100% / 100% / 100%', '行业由公司名+岗位名推断，发布时间按近12个月分布'],
        ['resume_clean.csv', 'school / major / projects / target_jobs', '100% / 100% / 100% / 100%', 'GitHub用户已补全教育背景，synthetic按JD技能分布生成'],
        ['skill_ontology.json', 'category / lifecycle_stage / parent_skill', '100% / 100% / 31%', '10个分类，3个生命周期阶段，169个有父技能'],
        ['github_detail.jsonl', 'language / stars / readme / skill_standard', '100% / 100% / 100% / 99%', '已补全star/fork/readme，language修正'],
        ['blog_trend.jsonl', 'summary / tags / hot_score', '100% / 100% / 100%', '已生成中文摘要和变化热度值'],
    ]
)

doc.add_page_break()

# ============================================================
# 第三章：知识图谱总体架构
# ============================================================
doc.add_heading('三、知识图谱总体架构', level=1)

doc.add_heading('3.1 设计原则', level=2)
principles = [
    ('全中文图数据库', '节点标签、关系类型、属性名全部使用中文，确保 Neo4j Browser 可视化展示和查询的可读性。'),
    ('无空格命名', '属性名统一不使用空格，采用中文简写，如"岗位ID""技能标准名""证据可信度"。'),
    ('数据分层与来源标注', '所有节点和关系保留 source_type 字段，区分 real_crawled / public_api / processed / synthetic / template_generated / manual_reviewed / algorithm。'),
    ('证据可追溯', '核心关系附带证据字段（source_url、evidence_score），支持 GraphRAG 回溯。'),
    ('动态预留', '通过 time_slice、from_slice、to_slice、change_type 字段为后续演化分析预留空间。'),
    ('LLM协同', 'LLM 生成岗位定义、学习路径、能力变化解释时，必须引用图谱和文本证据。'),
]
add_table(['原则', '说明'], principles)

doc.add_heading('3.2 分层架构', level=2)
add_para('知识图谱分为四层，按优先级递进构建：')
add_table(
    ['层级', '包含内容', '支撑功能', '优先级'],
    [
        ['基础图谱', '岗位、技能、人才、公司、课程、证书、技能分类；8种核心关系', '岗位查询、技能浏览、人才画像、基础匹配', 'P0 第一阶段'],
        ['增强图谱', '技术项目、论文、技术文章、证据节点；6种证据关系', 'GraphRAG问答、推荐解释、多源证据', 'P1 第二阶段'],
        ['演化图谱', '时间切片、能力变更事件、岗位定义版本、新岗位候选', '新岗位发现、既有岗位动态演化', 'P1 第二阶段'],
        ['评估图谱', '评估样本、系统评估关系、质量指标', '准确率评估、F1计算、异常检测', 'P2 第三阶段'],
    ]
)

doc.add_heading('3.3 节点-关系全景', level=2)
add_para('核心设计理念：以"技能"为中枢节点，岗位、人才、课程、证书、项目均通过技能建立关联，形成星型+网状混合拓扑。')
add_para('岗位 ←→ 技能 ←→ 人才')
add_para('              ←→ 课程/证书')
add_para('              ←→ 技术项目/论文/文章')

doc.add_page_break()

# ============================================================
# 第四章：核心节点设计
# ============================================================
doc.add_heading('四、核心节点设计（12类）', level=1)

add_para('所有节点使用中文标签，属性名使用中文简写。每个节点必须包含 source_type 和 evidence_score 字段。')

doc.add_heading('4.1 基础图谱节点（第一层）', level=2)

# 岗位节点
add_para('【节点1】岗位 (Job)', bold=True)
add_table(
    ['属性名', '类型', '必填', '说明', '示例值'],
    [
        ['岗位ID', 'String', '✓', '唯一标识，格式 JOB_00001', 'JOB_00001'],
        ['原始岗位名', 'String', '✓', '爬虫采集的原始名称', 'Java开发工程师（派遣编制）'],
        ['标准岗位名', 'String', '✓', '标准化后的名称', 'Java开发工程师'],
        ['公司名', 'String', '✓', '招聘公司', '中图数字科技(北京)有限公司'],
        ['行业', 'String', '✓', '行业分类', '信息技术'],
        ['工作地点', 'String', '', '城市', '北京'],
        ['薪资', 'String', '', '薪资范围', '1.8-2.3万'],
        ['学历要求', 'String', '', '学历', '本科'],
        ['经验要求', 'String', '', '工作年限', '5-10年'],
        ['岗位描述', 'String', '✓', '岗位概述', '负责Java后端服务的架构设计与开发...'],
        ['任职要求', 'String', '✓', '具体要求', '精通Java、Spring、MySQL...'],
        ['发布时间', 'String', '✓', '发布日期 YYYY-MM-DD', '2026-05-27'],
        ['来源URL', 'String', '✓', '原始链接', 'https://...'],
        ['来源名称', 'String', '✓', '数据来源平台', '前程无忧'],
        ['时间切片', 'String', '✓', '所属季度', '2026Q2'],
        ['来源类型', 'String', '✓', 'real_crawled/processed', 'real_crawled'],
        ['证据可信度', 'Float', '✓', '0.0~1.0', '0.88'],
        ['去重分数', 'Float', '✓', '重复度', '0.35'],
    ]
)

# 技能节点
add_para('【节点2】技能 (Skill)', bold=True)
add_table(
    ['属性名', '类型', '必填', '说明', '示例值'],
    [
        ['技能ID', 'String', '✓', '唯一标识 SKILL_00001', 'SKILL_00001'],
        ['技能标准名', 'String', '✓', '标准中文名', 'Python'],
        ['别名列表', 'List', '✓', '同义词', '["python","python3"]'],
        ['技能分类', 'String', '✓', '所属类别', 'AI'],
        ['父技能', 'String', '', '上级技能', '机器学习'],
        ['生命周期', 'String', '✓', 'mature/growth/emerging', 'mature'],
        ['来源类型', 'String', '✓', 'processed/manual', 'processed'],
    ]
)

# 技能分类节点
add_para('【节点3】技能分类 (SkillCategory)', bold=True)
add_table(
    ['属性名', '类型', '必填', '说明', '示例值'],
    [
        ['分类ID', 'String', '✓', '唯一标识', 'CAT_AI'],
        ['分类名', 'String', '✓', '分类名称', 'AI'],
        ['分类描述', 'String', '', '描述', '人工智能相关技能'],
    ]
)

# 人才节点
add_para('【节点4】人才 (Talent)', bold=True)
add_table(
    ['属性名', '类型', '必填', '说明', '示例值'],
    [
        ['人才ID', 'String', '✓', '唯一标识', 'gh_walidshaari'],
        ['人才类型', 'String', '✓', 'github_user/synthetic', 'github_user'],
        ['学历', 'String', '', '最高学历', '硕士'],
        ['专业', 'String', '', '所学专业', '计算机科学与技术'],
        ['学校', 'String', '', '毕业院校', '浙江大学'],
        ['所在地', 'String', '', '所在城市', '北京'],
        ['个人简介', 'String', '', '自我介绍', '硕士毕业于浙江大学...'],
        ['来源', 'String', '✓', '数据来源', 'github_public'],
        ['来源类型', 'String', '✓', '', 'public_api'],
    ]
)

# 课程节点
add_para('【节点5】课程 (Course)', bold=True)
add_table(
    ['属性名', '类型', '必填', '说明', '示例值'],
    [
        ['课程ID', 'String', '✓', '唯一标识', 'COURSE_00001'],
        ['课程名', 'String', '✓', '课程名称', '机器学习'],
        ['提供方', 'String', '✓', '平台/机构', 'Coursera'],
        ['难度', 'String', '✓', '初级/中级/高级', '中级'],
        ['时长', 'String', '', '学习时长', '40小时'],
        ['来源类型', 'String', '✓', 'template_generated', 'template_generated'],
    ]
)

# 证书节点
add_para('【节点6】证书 (Certificate)', bold=True)
add_table(
    ['属性名', '类型', '必填', '说明', '示例值'],
    [
        ['证书ID', 'String', '✓', '唯一标识', 'CERT_00001'],
        ['证书名', 'String', '✓', '证书名称', 'CKA认证'],
        ['颁发机构', 'String', '✓', '颁发方', 'CNCF'],
        ['等级', 'String', '✓', '初级/中级/高级', '高级'],
        ['来源类型', 'String', '✓', 'template_generated', 'template_generated'],
    ]
)

# 公司节点
add_para('【节点7】公司 (Company)', bold=True)
add_table(
    ['属性名', '类型', '必填', '说明', '示例值'],
    [
        ['公司ID', 'String', '✓', 'MD5 hash', 'CO_abc123'],
        ['公司名', 'String', '✓', '公司名称', '中图数字科技(北京)有限公司'],
        ['行业', 'String', '', '行业', '信息技术'],
    ]
)

# 项目节点
add_para('【节点8】项目 (Project)', bold=True)
add_table(
    ['属性名', '类型', '必填', '说明', '示例值'],
    [
        ['项目ID', 'String', '✓', '唯一标识', 'PROJ_00001'],
        ['项目名', 'String', '✓', '项目名称', '基于Python的智能推荐系统'],
        ['项目描述', 'String', '', '项目简介', '负责核心功能的设计与实现'],
        ['来源类型', 'String', '✓', '', 'synthetic'],
    ]
)

doc.add_heading('4.2 增强图谱节点（第二层）', level=2)

add_para('【节点9】技术项目 (TechProject) — GitHub开源仓库', bold=True)
add_para('【节点10】论文 (Paper) — arXiv学术论文')
add_para('【节点11】技术文章 (Blog) — 技术博客/行业报告')
add_para('【节点12】证据 (Evidence) — 文本证据元数据，关联向量库ID')

doc.add_heading('4.3 演化图谱节点（第三层）', level=2)
add_para('【节点13】岗位群 (JobCluster) — 同类岗位聚合')
add_para('【节点14】新岗位候选 (NewJobCandidate) — 发现的新兴岗位')
add_para('【节点15】岗位定义版本 (JobDefVersion) — 岗位定义的历史版本')
add_para('【节点16】时间切片 (TimeSlice) — 时间维度标识')
add_para('【节点17】能力变更事件 (AbilityChange) — 技能变化记录')

doc.add_page_break()

# ============================================================
# 第五章：核心关系设计
# ============================================================
doc.add_heading('五、核心关系设计（18种）', level=1)

add_para('所有关系使用中文类型名，附带权重、证据可信度、来源类型等属性。', bold=True)

doc.add_heading('5.1 基础图谱关系（第一层，8种）', level=2)

add_table(
    ['关系类型', '起点', '终点', '关键属性', '数据来源'],
    [
        ['发布岗位', '公司', '岗位', 'source_url, time_slice', 'jd_clean.csv'],
        ['要求技能', '岗位', '技能', '权重, 要求类型(必备/加分), 证据可信度', 'jd_clean skill_standard'],
        ['属于分类', '技能', '技能分类', '置信度', 'skill_ontology category'],
        ['父技能', '技能', '技能', '置信度, 层级', 'skill_ontology parent_skill'],
        ['拥有技能', '人才', '技能', '掌握程度, 可信度, 来源类型', 'resume_clean skill_standard'],
        ['参与项目', '人才', '项目', '角色, 时间', 'resume_clean projects'],
        ['使用技能', '项目', '技能', '置信度', 'resume_clean / github_detail'],
        ['教授技能', '课程', '技能', '难度, 时长, 来源类型', 'course_data skills'],
    ]
)

doc.add_heading('5.2 增强图谱关系（第二层，6种）', level=2)

add_table(
    ['关系类型', '起点', '终点', '关键属性', '数据来源'],
    [
        ['认证技能', '证书', '技能', '等级, 颁发机构', 'certificate_data'],
        ['使用技术', '技术项目', '技能', '热度, stars', 'github_detail skill_standard'],
        ['涉及技术', '论文', '技能', '发布时间, 热度', 'arxiv_trend'],
        ['涉及技术', '技术文章', '技能', '发布时间, 热度', 'blog_trend'],
        ['有证据', '各类节点', '证据', '证据可信度, vector_id', '全源'],
        ['人岗匹配', '人才', '岗位', '匹配来源, 匹配等级, 匹配得分, 匹配技能, 缺失技能', 'algorithm / manual_reviewed'],
    ]
)

doc.add_heading('5.3 演化图谱关系（第三层，4种）', level=2)

add_table(
    ['关系类型', '起点', '终点', '关键属性', '说明'],
    [
        ['属于岗位群', '岗位', '岗位群', '置信度, 规则来源', '标准岗位名聚合'],
        ['聚合岗位', '岗位群', '岗位', '来源数量, 时间切片', '逆向关系'],
        ['演化', '岗位群', '能力变更事件', '变化类型, 时间切片, 证据数量', '多期对比'],
        ['涉及技能', '能力变更事件', '技能', '变化角色(新增/删除/强化/弱化)', '变化记录'],
    ]
)

doc.add_page_break()

# ============================================================
# 第六章：Neo4j 建图规范
# ============================================================
doc.add_heading('六、Neo4j 建图规范与 Cypher 示例', level=1)

doc.add_heading('6.1 命名规范', level=2)
add_table(
    ['元素', '规范', '示例'],
    [
        ['节点标签', '使用中文，反引号包裹', '`岗位`, `技能`, `人才`'],
        ['关系类型', '使用中文，反引号包裹', '`要求技能`, `拥有技能`'],
        ['属性名', '使用中文简写，无空格', '`岗位ID`, `技能标准名`, `证据可信度`'],
        ['约束名', 'constraint_ + 属性名 + _unique', 'constraint_岗位ID_unique'],
        ['索引名', 'index_ + 标签 + _ + 属性', 'index_岗位_标准岗位名'],
    ]
)

doc.add_heading('6.2 约束与索引', level=2)
add_para('-- 唯一性约束', bold=True)
cypher_examples = [
    "CREATE CONSTRAINT constraint_岗位ID_unique IF NOT EXISTS FOR (n:`岗位`) REQUIRE n.`岗位ID` IS UNIQUE;",
    "CREATE CONSTRAINT constraint_技能ID_unique IF NOT EXISTS FOR (n:`技能`) REQUIRE n.`技能ID` IS UNIQUE;",
    "CREATE CONSTRAINT constraint_人才ID_unique IF NOT EXISTS FOR (n:`人才`) REQUIRE n.`人才ID` IS UNIQUE;",
    "CREATE CONSTRAINT constraint_课程ID_unique IF NOT EXISTS FOR (n:`课程`) REQUIRE n.`课程ID` IS UNIQUE;",
    "CREATE CONSTRAINT constraint_证书ID_unique IF NOT EXISTS FOR (n:`证书`) REQUIRE n.`证书ID` IS UNIQUE;",
]
for c in cypher_examples:
    add_para(c, size=9)

doc.add_heading('6.3 核心查询示例', level=2)

add_para('【查询1】岗位核心技能及证据', bold=True)
add_para("""MATCH (j:`岗位`)-[r:`要求技能`]->(s:`技能`)
WHERE j.`标准岗位名` = "大模型应用开发工程师"
OPTIONAL MATCH (j)-[:`有证据`]->(e:`证据`)
RETURN j.`原始岗位名`, s.`技能标准名`, r.`权重`, r.`证据可信度`, e.`来源URL`
ORDER BY r.`证据可信度` DESC LIMIT 50;""", size=9)

add_para('【查询2】人才与目标岗位的能力差距', bold=True)
add_para("""MATCH (t:`人才` {`人才ID`: $talent_id})-[:`拥有技能`]->(ts:`技能`)
WITH t, collect(ts.`技能标准名`) AS talent_skills
MATCH (j:`岗位` {`标准岗位名`: $target_job})-[:`要求技能`]->(js:`技能`)
WITH t, talent_skills, collect(js.`技能标准名`) AS job_skills
RETURN [x IN job_skills WHERE x IN talent_skills] AS 匹配技能,
       [x IN job_skills WHERE NOT x IN talent_skills] AS 缺失技能;""", size=9)

add_para('【查询3】技术趋势证据（多源交叉验证）', bold=True)
add_para("""MATCH (tp:`技术项目`)-[:`使用技术`]->(s:`技能`)
WHERE s.`生命周期` IN ["growth", "emerging"]
RETURN s.`技能标准名`, count(tp) AS 项目证据数, avg(tp.`热度`) AS 平均热度
ORDER BY 项目证据数 DESC LIMIT 20;""", size=9)

add_para('【查询4】岗位能力随时间的变化', bold=True)
add_para("""MATCH (j:`岗位`)-[:`要求技能`]->(s:`技能`)
WHERE j.`行业` = "人工智能"
WITH j.`时间切片` AS 切片, s.`技能标准名` AS 技能, count(*) AS 频率
ORDER BY 切片, 频率 DESC
RETURN 切片, collect({技能: 技能, 频率: 频率})[0..10] AS Top10技能;""", size=9)

doc.add_page_break()

# ============================================================
# 第七章：数据到图谱的映射
# ============================================================
doc.add_heading('七、数据到图谱的映射流程', level=1)

doc.add_heading('7.1 ETL 映射总表', level=2)
add_table(
    ['数据文件', '生成节点', '生成关系', '处理规则'],
    [
        ['jd_clean.csv', '岗位、公司', '公司-发布岗位→岗位, 岗位-要求技能→技能', '每行→1个岗位节点, 公司去重, skill_standard分号分隔→多个关系'],
        ['resume_clean.csv', '人才、项目', '人才-拥有技能→技能, 人才-参与项目→项目, 项目-使用技能→技能', '每行→1个人才节点, projects JSON解析→项目'],
        ['skill_ontology.json', '技能、技能分类', '技能-属于分类→技能分类, 技能-父技能→技能', '每个key→1个技能节点, category去重→分类节点'],
        ['job_standard_dict.csv', '岗位群', '岗位-属于岗位群→岗位群', '标准岗位名去重→岗位群'],
        ['course_data.csv', '课程', '课程-教授技能→技能', 'skills分号分隔→关系'],
        ['certificate_data.csv', '证书', '证书-认证技能→技能', 'related_skills→关系'],
        ['github_detail.jsonl', '技术项目', '技术项目-使用技术→技能', 'skill_standard→关系, 附加热度/star'],
        ['arxiv_trend.jsonl', '论文', '论文-涉及技术→技能', 'tech_name+tags→技能关系'],
        ['blog_trend.jsonl', '技术文章', '技术文章-涉及技术→技能', 'tech_name+tags→技能关系'],
    ]
)

doc.add_heading('7.2 导入流程', level=2)
add_para('步骤1：生成节点CSV — 从各数据源提取节点属性，写入 nodes_*.csv')
add_para('步骤2：生成关系CSV — 建立节点间关系，写入 rel_*.csv')
add_para('步骤3：Neo4j Admin Import — 使用 neo4j-admin import 批量导入')
add_para('步骤4：创建约束和索引 — 执行 Cypher 约束语句')
add_para('步骤5：数据质量校验 — 检查孤立节点、关系完整性、source_type覆盖')

doc.add_page_break()

# ============================================================
# 第八章：新岗位发现算法
# ============================================================
doc.add_heading('八、新岗位发现算法设计', level=1)

doc.add_heading('8.1 发现流程', level=2)
add_para('新岗位发现采用"数据驱动 + 模型生成 + 证据校验 + 人工确认"四步法：')
add_para('Step 1 — 技能共现聚类：从JD中提取技能共现矩阵，使用社区发现算法识别新兴技能组合')
add_para('Step 2 — 岗位簇识别：将技能组合与现有岗位群对比，识别未命中现有岗位定义的新簇')
add_para('Step 3 — LLM生成定义：将新簇的JD片段和技能列表发送给LLM，生成岗位名称、职责、必备/加分技能、典型行业')
add_para('Step 4 — 证据校验：用GitHub/arXiv/Blog数据交叉验证技能热度，计算稳定性评分')
add_para('Step 5 — 人工确认：管理员审核确认后，创建"新岗位候选"节点并入图')

doc.add_heading('8.2 评分公式', level=2)
add_para('new_job_score = 岗位簇规模 × 0.35 + 新兴技能强度 × 0.25 + 行业多样性 × 0.20 + 外部证据支持度 × 0.10 + 人工确认权重 × 0.10')
add_para('阈值：score ≥ 0.6 → 高置信度新岗位；0.4~0.6 → 待观察；< 0.4 → 暂不纳入')

doc.add_heading('8.3 示例', level=2)
add_table(
    ['候选岗位', '核心技能', '证据来源', '稳定性评分', '状态'],
    [
        ['AI Agent应用工程师', 'Python, Agent, RAG, MCP, LangChain, 向量数据库', 'JD 48条 + GitHub 29个仓库 + 论文 15篇', '0.82', '✅ 已确认'],
        ['RAG系统工程师', 'Python, RAG, 向量数据库, LLM, LangChain, Milvus', 'JD 35条 + GitHub 25个仓库', '0.75', '✅ 已确认'],
        ['提示词工程师', 'Prompt Engineering, LLM, NLP', 'JD 12条', '0.45', '⏳ 待观察'],
    ]
)

doc.add_page_break()

# ============================================================
# 第九章：动态演化
# ============================================================
doc.add_heading('九、既有岗位能力动态演化', level=1)

doc.add_heading('9.1 演化检测机制', level=2)
add_para('通过时间切片（季度）对比同一岗位群在不同时期的能力要求变化：')
add_para('① 新增技能：当前切片高频出现、历史切片未出现')
add_para('② 删除/淘汰技能：历史切片高频出现、当前切片几乎消失')
add_para('③ 强化技能：从"加分"变为"必备"')
add_para('④ 弱化技能：从"必备"变为"加分"')
add_para('⑤ 技能合并/拆分：多个旧技能合并为一个新综合技能，或一个技能拆分为多个子技能')

doc.add_heading('9.2 变化检测阈值', level=2)
add_table(
    ['变化类型', '判定条件', '图操作'],
    [
        ['新增/强化', '当前切片出现频率 > 历史均值 × 1.5 且达到最小支持度(≥3)', '创建或强化"要求技能"关系，权重+1'],
        ['弱化', '当前频率 < 历史均值 × 0.5 且下降趋势持续≥2个切片', '关系标记为 weakened'],
        ['删除/淘汰', '连续2个切片频率为0', '关系标记为 expired，保留历史关系'],
        ['修改', 'required↔bonus 类型变化 或 标准技能名替换', '记录 before/after 到能力变更事件'],
    ]
)

doc.add_heading('9.3 演化示例', level=2)
add_para('以"Java开发工程师"为例，2025Q3→2026Q2的技能要求变化：')
add_table(
    ['技能', '2025Q3', '2026Q2', '变化'],
    [
        ['Java', '必备(100%)', '必备(98%)', '无变化 — 核心稳定'],
        ['Spring Boot', '必备(85%)', '必备(92%)', '↑ 强化 — 持续核心'],
        ['Docker', '加分(45%)', '必备(72%)', '↑ 升级 — 从加分变为必备'],
        ['大语言模型', '未出现(0%)', '加分(18%)', '↑ 新增 — AI技能向传统后端渗透'],
        ['SOAP WebService', '加分(35%)', '加分(5%)', '↓ 弱化 — 逐步被REST/gRPC替代'],
        ['Struts', '加分(22%)', '未出现(0%)', '→ 淘汰 — 历史框架退出市场'],
    ]
)

doc.add_page_break()

# ============================================================
# 第十章：人岗匹配
# ============================================================
doc.add_heading('十、人岗匹配与能力差距分析', level=1)

doc.add_heading('10.1 三维匹配模型', level=2)
add_table(
    ['维度', '权重', '计算方式'],
    [
        ['技能匹配 (40%)', '40%', '人才技能 ∩ 岗位必备技能 / 岗位必备技能总数'],
        ['项目匹配 (30%)', '30%', '项目技术栈与岗位技能要求的重叠度'],
        ['发展潜力 (30%)', '30%', '专业背景+相关技能+学习能力的综合评估'],
    ]
)
add_para('综合匹配度 = 技能匹配 × 0.40 + 项目匹配 × 0.30 + 发展潜力 × 0.30')
add_para('匹配等级：高匹配(≥80%) / 中匹配(60-79%) / 低匹配(40-59%) / 不匹配(<40%)')

doc.add_heading('10.2 能力差距分析', level=2)
add_para('差距 = 目标岗位核心技能集合 - 候选人已有技能集合')
add_para('差距按优先级排序：必备技能差距 > 加分技能差距 > 进阶技能差距')
add_para('每个差距技能附带：学习难度(易/中/难)、预计学习周期、推荐课程/证书/项目')

doc.add_page_break()

# ============================================================
# 第十一章：GraphRAG
# ============================================================
doc.add_heading('十一、GraphRAG 证据链与防幻觉设计', level=1)

doc.add_heading('11.1 双层检索架构', level=2)
add_para('GraphRAG = Neo4j 图路径检索（结构化） + Qdrant 文本向量检索（非结构化） + LLM 整合生成')
add_para('① 图路径检索：从 Neo4j 中检索相关岗位、技能、课程、项目的关联路径')
add_para('② 文本向量检索：从 Qdrant 中检索 JD片段、GitHub README、论文摘要、博客摘要')
add_para('③ LLM整合：将图路径作为结构证据，文本片段作为内容证据，共同输入LLM生成回答')

doc.add_heading('11.2 防幻觉机制', level=2)
add_table(
    ['策略', '实现方式'],
    [
        ['证据强制', 'Prompt要求所有结论必须附带证据来源，无证据不生成结论'],
        ['低分过滤', 'evidence_score < 阈值(0.3) 的节点/关系不参与检索'],
        ['来源可追溯', '前端展示：回答 → 引用来源URL → 图谱路径 → 文本片段'],
        ['冲突标注', '不同来源对同一变化判断冲突时，标记为 conflict 并推送人工审核'],
        ['置信度分级', '回答按证据数量和质量分：高置信(≥3源) / 中置信(2源) / 低置信(1源)'],
    ]
)

doc.add_heading('11.3 应用场景', level=2)
add_table(
    ['场景', '问题示例', '使用证据'],
    [
        ['岗位定义查询', '"大模型应用工程师需要什么技能？"', 'JD图谱路径 + JD原文片段'],
        ['能力变化解释', '"为什么Docker从加分变成必备？"', '时间切片对比 + 行业报告摘录'],
        ['匹配推荐解释', '"为什么推荐我投AI Agent岗位？"', '技能匹配图谱 + 项目经验片段'],
        ['学习路径推荐', '"我想转AI方向，应该学什么？"', '能力差距图谱 + 课程描述片段'],
    ]
)

doc.add_page_break()

# ============================================================
# 第十二章：学习路径
# ============================================================
doc.add_heading('十二、学习路径推荐', level=1)

add_para('学习路径的生成逻辑：')
add_para('Step 1：从"人岗匹配"模块获取能力差距列表')
add_para('Step 2：按优先级排序（必备技能差距优先）')
add_para('Step 3：对每个差距技能，查询课程-教授技能和证书-认证技能关系，召回学习资源')
add_para('Step 4：查询技术项目-使用技术关系，推荐实战项目')
add_para('Step 5：LLM 按"基础→核心→实战→进阶→检验"五阶梯排序，生成学习路径')
add_para('Step 6：每条学习任务标注：学习资源、预计时间、检验标准')

doc.add_page_break()

# ============================================================
# 第十三章：评估指标
# ============================================================
doc.add_heading('十三、评估指标体系', level=1)

add_para('基于赛题要求的 300 条金标数据集（Gold JD 100 + Gold Resume 50 + Match Label 100 + Negative 50），系统需达到以下指标：')

add_table(
    ['指标', '赛题要求', '评估方式', '金标数据'],
    [
        ['JD技能提取准确率 (Precision)', '≥90%', '对比 gold_jd_set_reviewed 中人工标注的 required/bonus 技能', '100条金标JD'],
        ['JD技能提取召回率 (Recall)', '≥90%', '系统提取技能 ∩ 人工标注技能 / 人工标注技能', '100条金标JD'],
        ['简历要素提取准确率', '≥90%', '对比 gold_resume_set_reviewed 中人工标注的技能/项目/学历', '50条金标简历'],
        ['人岗匹配准确率', '≥90%', '对比 match_label_set_reviewed 中人工标注的 high/medium/low', '100条金标匹配'],
        ['异常检测召回率', '≥85%', '对比 negative_samples_reviewed 中的人工标注', '50条负样本'],
        ['GraphRAG证据命中率', '≥90%', '有效证据的回答数 / 总回答数', '人工评估'],
    ]
)

doc.add_page_break()

# ============================================================
# 第十四章：实施路线图
# ============================================================
doc.add_heading('十四、分阶段实施路线图', level=1)

doc.add_heading('14.1 第一阶段：基础图谱 + 核心闭环（2-3周）', level=2)
add_para('目标：跑通"数据→图谱→查询→前端展示"全流程')
add_para('节点：岗位、技能、技能分类、人才、公司、课程、证书、项目（8类）')
add_para('关系：发布岗位、要求技能、属于分类、拥有技能、参与项目、使用技能、教授技能、认证技能（8种）')
add_para('前端：管理员端首页看板、新岗位发现、知识图谱可视化；学生端简历上传、岗位推荐（数据对接真实图谱）')

doc.add_heading('14.2 第二阶段：增强图谱 + GraphRAG（3-4周）', level=2)
add_para('目标：引入多源证据和智能问答')
add_para('节点：技术项目、论文、技术文章、证据、新岗位候选、能力变更事件（6类）')
add_para('关系：使用技术、涉及技术、有证据、人岗匹配、属于岗位群、演化、涉及技能（7种）')
add_para('新增：Qdrant向量库 + GraphRAG证据链 + 动态演化检测 + 新岗位发现')

doc.add_heading('14.3 第三阶段：评估 + 优化（1-2周）', level=2)
add_para('目标：质量达标和体验优化')
add_para('节点：评估样本、学习路径、学习路径步骤（3类）')
add_para('关系：学习路径-包含步骤、评估样本-标注等')
add_para('新增：金标评估仪表盘、学习路径自动生成、系统性能优化')

doc.add_page_break()

# ============================================================
# 附录
# ============================================================
doc.add_heading('附录A：节点属性完整清单', level=1)
add_para('详见第四章，共 12 类核心节点、5 类演化节点。每类节点的必填属性汇总如下：')
add_para('所有节点共有属性：source_type(来源类型), evidence_score(证据可信度)')
add_para('岗位节点必填：岗位ID, 原始岗位名, 标准岗位名, 公司名, 行业, 岗位描述, 任职要求, 发布时间, 来源URL, 时间切片')
add_para('技能节点必填：技能ID, 技能标准名, 别名列表, 技能分类, 生命周期')
add_para('人才节点必填：人才ID, 人才类型, 来源')
add_para('关系共有属性：权重(部分), 证据可信度, 来源类型, 时间切片(部分)')

doc.add_heading('附录B：Cypher 查询参考', level=1)
add_para('详见第六章 6.3 节，包含4个核心查询模板：')
add_para('1. 岗位核心技能及证据查询')
add_para('2. 人才能力差距分析查询')
add_para('3. 技术趋势多源证据查询')
add_para('4. 岗位能力时间演化查询')

doc.add_heading('附录C：数据文件清单', level=1)
add_table(
    ['序号', '文件', '规模', '用途'],
    [
        ['1', 'clean/jd_clean.csv', '2,500条', '岗位+公司节点'],
        ['2', 'clean/resume_clean.csv', '500条', '人才+项目节点'],
        ['3', 'meta/skill_ontology.json', '545个', '技能+分类节点'],
        ['4', 'meta/job_standard_dict.csv', '37条', '岗位群节点'],
        ['5', 'raw/github_detail.jsonl', '389条', '技术项目节点'],
        ['6', 'raw/github_trend.jsonl', '389条', '技术趋势数据'],
        ['7', 'raw/arxiv_trend.jsonl', '359条', '论文节点'],
        ['8', 'raw/blog_trend.jsonl', '857条', '技术文章节点'],
        ['9', 'education/course_data.csv', '384条', '课程节点'],
        ['10', 'education/certificate_data.csv', '101条', '证书节点'],
        ['11', 'meta/gold_jd_set_reviewed.json', '100条', 'JD评估'],
        ['12', 'meta/gold_resume_set_reviewed.json', '50条', '简历评估'],
        ['13', 'meta/match_label_set_reviewed.json', '100条', '匹配评估'],
        ['14', 'meta/negative_samples_reviewed.json', '50条', '异常评估'],
    ]
)

# ============================================================
# 结尾
# ============================================================
doc.add_paragraph()
add_para('— 文档结束 —', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('TalentGraph Evolution 知识图谱设计文档 V4.0', align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('基于赛题 XH-202621 · 中文知识图谱 · Neo4j · GraphRAG', align=WD_ALIGN_PARAGRAPH.CENTER)

# 保存
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
doc.save(OUTPUT)
print(f'文档已生成: {OUTPUT}')
