#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成中文知识图谱详细设计文档 Word 版，所有节点/关系/属性均使用中文命名，适配 Neo4j 浏览器中文展示。"""
import json, os, sys
from collections import Counter
from datetime import datetime

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, BASE)
ROOT = os.path.dirname(BASE) if os.path.basename(BASE) == 'crawler' else os.path.dirname(BASE)

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def load_ontology():
    p = os.path.join(BASE, 'data', 'meta', 'skill_ontology.json')
    with open(p, encoding='utf-8') as f:
        return json.load(f)

def shade(cell, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color)
    s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(c, '2B579A')
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i + 1].cells[j]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if i % 2 == 1:
                shade(c, 'F2F2F2')
    doc.add_paragraph()
    return t

def heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def para(doc, text):
    doc.add_paragraph(text)

def bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')

def code_block(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(1)

# ================================================================
def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    for m in [sec.left_margin, sec.right_margin, sec.top_margin, sec.bottom_margin]:
        m.value = int(Cm(2.5))

    # 封面
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('多源异构数据驱动岗位和能力图谱\n知识图谱详细设计文档')
    r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    doc.add_paragraph()
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('Neo4j 中文知识图谱 — 全中文节点/关系/属性设计')
    r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph(); doc.add_paragraph()
    for line in [
        '文档版本: V2.0 (中文图谱版)',
        f'编制日期: {datetime.now().strftime("%Y年%m月%d日")}',
        '所属项目: TalentGraph Evolution',
        '比赛题目: XH-202621 多源异构数据驱动岗位和能力图谱构建与动态演化分析研究',
        '发榜单位: 科大讯飞股份有限公司',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line); r.font.size = Pt(11)
    doc.add_page_break()

    # ================================================================
    heading(doc, '一、项目概述')
    heading(doc, '1.1 项目背景', 2)
    para(doc,
        '随着数字经济的快速发展，以人工智能、大数据、云计算为代表的新一代信息技术'
        '正成为产业升级的核心引擎。技术迭代速度远超人才培养周期，劳动力市场出现'
        '结构性矛盾——企业在新兴岗位上面临招不到合适的人的困境。'
        '本赛题（XH-202621）由科大讯飞发榜，聚焦[数据驱动+大模型+知识图谱]'
        '核心技术方向，要求构建能够自我进化的人才能力大脑。'
    )

    heading(doc, '1.2 Neo4j 中文图谱设计说明', 2)
    para(doc,
        '本知识图谱全部采用中文命名体系，确保在 Neo4j Browser 中直接呈现中文可读的'
        '节点标签、关系类型和属性名称。具体约定如下：\n'
        '  - 节点标签（Label）：全部使用中文，如 岗位、技能、人才、企业\n'
        '  - 关系类型（Relationship Type）：全部使用中文，如 要求技能、拥有技能、人岗匹配\n'
        '  - 属性键（Property Key）：全部使用中文，如 岗位名称、技能标准名、学历要求\n'
        '  - 属性值：保留原始中文内容\n'
        '这样在 Neo4j Browser 中执行 MATCH (n:岗位) RETURN n 时，'
        '节点直接以 岗位 标签展示，属性以中文键名显示，可视化效果直观可读。'
    )

    heading(doc, '1.3 核心目标', 2)
    for g in [
        '构建覆盖新一代信息技术领域的岗位能力知识图谱（全中文），颗粒度到技能点级别',
        '实现新岗位的自动发现与定义，支持人工优化与动态更新',
        '识别既有岗位能力要求的变化，标注新增、删除、修改的技能项',
        '实现高精度简历解析与人岗匹配差距分析，提供针对性改进建议和学习路径',
    ]:
        bullet(doc, g)

    doc.add_page_break()

    # ================================================================
    heading(doc, '二、中文节点设计')

    heading(doc, '2.1 节点总览（9类）', 2)
    para(doc,
        '知识图谱共设计9类节点，所有节点标签（Label）使用中文。'
        '在 Neo4j 中创建节点时，标签名即为下表的中文标签。'
    )
    table(doc,
        ['序号', 'Neo4j 节点标签', '中文含义', '节点标识属性', '数据来源', '预估数量'],
        [
            ['1', '岗位', '招聘岗位', '岗位ID', 'jd_clean.csv + gold_jd_reviewed.json', '2,500'],
            ['2', '技能', '技术技能', '技能标准名', 'skill_ontology.json', '545'],
            ['3', '人才', '求职者/候选人', '人才ID', 'resume_clean.csv + gold_resume_reviewed.json', '180'],
            ['4', '企业', '招聘企业', '企业名称', 'jd_clean.csv 去重', '~2,000'],
            ['5', '课程', '培训课程', '课程ID', 'course_data.csv', '384'],
            ['6', '证书', '职业认证', '证书ID', 'certificate_data.csv', '101'],
            ['7', '技术项目', '开源技术项目', '项目ID', 'github_trend.jsonl', '389'],
            ['8', '论文', '学术论文', '论文ID', 'arxiv_trend.jsonl', '359'],
            ['9', '技能分类', '技能所属技术方向', '分类名称', 'ontology category 去重', '10'],
        ],
    )

    heading(doc, '2.2 岗位 节点详细属性', 2)
    para(doc, 'Neo4j 标签: 岗位')
    table(doc,
        ['中文属性名', 'Neo4j 属性键', '类型', '示例值'],
        [
            ['岗位ID', '岗位ID', '字符串', 'JOB_00001'],
            ['原始岗位名', '原始岗位名', '字符串', '大模型应用开发工程师'],
            ['标准岗位名', '标准岗位名', '字符串', '大模型应用开发工程师'],
            ['企业名称', '企业名称', '字符串', '北京爱创科技股份有限公司'],
            ['工作地点', '工作地点', '字符串', '北京'],
            ['薪资范围', '薪资范围', '字符串', '2-2.5万'],
            ['学历要求', '学历要求', '字符串', '本科'],
            ['经验要求', '经验要求', '字符串', '5-10年'],
            ['岗位难度', '岗位难度', '字符串', '高级'],
            ['时间切片', '时间切片', '字符串', '2026Q3'],
            ['证据可信度', '证据可信度', '浮点数', '0.93'],
            ['数据来源', '数据来源', '字符串', '智联招聘'],
            ['来源链接', '来源链接', '字符串', 'http://www.zhaopin.com/...'],
        ],
    )

    heading(doc, '2.3 技能 节点详细属性', 2)
    para(doc, 'Neo4j 标签: 技能')
    table(doc,
        ['中文属性名', 'Neo4j 属性键', '类型', '示例值'],
        [
            ['技能标准名', '技能标准名', '字符串', '大语言模型'],
            ['别名列表', '别名列表', '字符串列表', 'LLM, 大模型, large language model'],
            ['所属分类', '所属分类', '字符串', 'AI'],
            ['父技能', '父技能', '字符串', '人工智能'],
            ['生命周期', '生命周期', '字符串', '增长期'],
            ['出现频次', '出现频次', '整数', '1250'],
        ],
    )

    heading(doc, '2.4 人才 节点详细属性', 2)
    para(doc, 'Neo4j 标签: 人才')
    table(doc,
        ['中文属性名', 'Neo4j 属性键', '类型', '示例值'],
        [
            ['人才ID', '人才ID', '字符串', 'TALENT_00001'],
            ['来源类型', '来源类型', '字符串', 'github用户'],
            ['学历', '学历', '字符串', '硕士'],
            ['毕业院校', '毕业院校', '字符串', '北京大学'],
            ['专业', '专业', '字符串', '软件工程'],
            ['目标岗位', '目标岗位', '字符串列表', 'AI工程师, 算法工程师'],
            ['数据类型', '数据类型', '字符串', '合成数据'],
        ],
    )

    heading(doc, '2.5 其余节点属性一览', 2)
    table(doc,
        ['节点标签', '核心属性（中文键名）'],
        [
            ['企业', '企业名称, 所属行业'],
            ['课程', '课程ID, 课程名称, 提供平台, 难度, 时长, 课程技能'],
            ['证书', '证书ID, 证书名称, 颁发机构, 等级, 关联技能'],
            ['技术项目', '项目ID, 项目名称, 摘要, 标签, 来源链接, 热度'],
            ['论文', '论文ID, 论文标题, 摘要, 标签, 来源链接, 发表时间'],
            ['技能分类', '分类名称'],
        ],
    )

    doc.add_page_break()

    # ================================================================
    heading(doc, '三、中文关系设计')

    heading(doc, '3.1 关系总览（11种）', 2)
    para(doc,
        '所有关系类型（Relationship Type）使用中文命名。'
        'Neo4j 中关系方向为有向箭头，查询时可用 -[:关系名]-> 语法。'
    )
    table(doc,
        ['序号', 'Neo4j 关系类型', '中文含义', '方向', '基数', '数据来源'],
        [
            ['1', '要求技能', '岗位必备技能', '岗位 -> 技能', '1:N', 'gold_jd的required_skills'],
            ['2', '加分技能', '岗位加分技能', '岗位 -> 技能', '1:N', 'gold_jd的bonus_skills'],
            ['3', '拥有技能', '人才掌握技能', '人才 -> 技能', '1:N', 'resume的skill_standard'],
            ['4', '属于分类', '技能归属方向', '技能 -> 技能分类', 'N:1', 'ontology的category'],
            ['5', '父技能', '技能层级归属', '技能 -> 技能', 'N:1', 'ontology的parent_skill'],
            ['6', '人岗匹配', '人才匹配岗位', '人才 -> 岗位', 'N:M', 'match_label_reviewed'],
            ['7', '教授技能', '课程涵盖技能', '课程 -> 技能', '1:N', 'course_data的skills'],
            ['8', '认证技能', '证书对应技能', '证书 -> 技能', '1:N', 'certificate_data的related_skills'],
            ['9', '发布岗位', '企业发布岗位', '企业 -> 岗位', '1:N', 'jd_clean的company'],
            ['10', '使用技术', '项目使用技术栈', '技术项目 -> 技能', '1:N', 'github_trend的tags'],
            ['11', '涉及技术', '论文涉及技术', '论文 -> 技能', '1:N', 'arxiv_trend的tags'],
        ],
    )

    heading(doc, '3.2 关系属性', 2)
    para(doc, '部分关系需要附加属性，用于精细化的查询和匹配计算：')
    table(doc,
        ['关系类型', '关系属性（中文键名）', '说明'],
        [
            ['要求技能', '权重', '0-1之间的数值，表示该技能对岗位的重要程度'],
            ['加分技能', '权重', '0-1之间的数值，通常低于要求技能'],
            ['拥有技能', '掌握程度, 置信度', '表示人才对该技能的掌握水平和识别置信度'],
            ['人岗匹配', '匹配等级, 匹配得分, 匹配时间', '等级: 高匹配/中匹配/低匹配；得分: 0-100'],
            ['使用技术', '频次', '该技术在项目中出现的次数或权重'],
        ],
    )

    doc.add_page_break()

    # ================================================================
    heading(doc, '四、Cypher 建图语句（中文版）')

    heading(doc, '4.1 创建约束与索引', 2)
    code_block(doc,
        '// 岗位约束\n'
        'CREATE CONSTRAINT 岗位ID约束 IF NOT EXISTS FOR (n:岗位) REQUIRE n.岗位ID IS UNIQUE;\n'
        '// 技能约束\n'
        'CREATE CONSTRAINT 技能名约束 IF NOT EXISTS FOR (n:技能) REQUIRE n.技能标准名 IS UNIQUE;\n'
        '// 人才约束\n'
        'CREATE CONSTRAINT 人才ID约束 IF NOT EXISTS FOR (n:人才) REQUIRE n.人才ID IS UNIQUE;\n'
        '// 企业约束\n'
        'CREATE CONSTRAINT 企业名约束 IF NOT EXISTS FOR (n:企业) REQUIRE n.企业名称 IS UNIQUE;\n'
        '// 技能分类约束\n'
        'CREATE CONSTRAINT 分类名约束 IF NOT EXISTS FOR (n:技能分类) REQUIRE n.分类名称 IS UNIQUE;'
    )

    heading(doc, '4.2 创建技能分类节点', 2)
    code_block(doc,
        'CREATE (:技能分类 {分类名称: "AI"});\n'
        'CREATE (:技能分类 {分类名称: "后端开发"});\n'
        'CREATE (:技能分类 {分类名称: "数据工程"});\n'
        'CREATE (:技能分类 {分类名称: "云计算"});\n'
        'CREATE (:技能分类 {分类名称: "前端开发"});\n'
        'CREATE (:技能分类 {分类名称: "AI智能体"});\n'
        'CREATE (:技能分类 {分类名称: "数据库"});\n'
        'CREATE (:技能分类 {分类名称: "物联网/嵌入式"});\n'
        'CREATE (:技能分类 {分类名称: "开发运维"});\n'
        'CREATE (:技能分类 {分类名称: "安全"});'
    )

    heading(doc, '4.3 批量导入技能节点', 2)
    para(doc,
        '使用 Neo4j 的 apoc.load.json 或 Python neo4j driver 批量导入。'
        '以下为 Python 导入示例：'
    )
    code_block(doc,
        'from neo4j import GraphDatabase\n'
        'driver = GraphDatabase.driver("bolt://localhost:7687", auth=("neo4j", "password"))\n'
        'with driver.session() as session:\n'
        '    for skill in skills_data:\n'
        '        session.run(\n'
        '            "MERGE (s:技能 {技能标准名: $name}) "\n'
        '            "SET s.别名列表 = $aliases, s.所属分类 = $cat, "\n'
        '            "    s.父技能 = $parent, s.生命周期 = $stage, s.出现频次 = $freq",\n'
        '            name=skill["standard_name"], aliases=skill["aliases"],\n'
        '            cat=skill["category"], parent=skill.get("parent_skill",""),\n'
        '            stage=skill["lifecycle_stage"], freq=skill.get("frequency",0)\n'
        '        )'
    )

    heading(doc, '4.4 创建关系示例', 2)
    para(doc, '以下是常用关系的 Cypher 创建语句：')
    code_block(doc,
        '// 技能属于分类\n'
        'MATCH (s:技能 {技能标准名: "Python"})\n'
        'MATCH (c:技能分类 {分类名称: "AI"})\n'
        'MERGE (s)-[:属于分类]->(c);\n\n'
        '// 岗位要求技能\n'
        'MATCH (j:岗位 {岗位ID: "JOB_00001"})\n'
        'MATCH (s:技能 {技能标准名: "Python"})\n'
        'MERGE (j)-[:要求技能 {权重: 0.9}]->(s);\n\n'
        '// 人才拥有技能\n'
        'MATCH (t:人才 {人才ID: "TALENT_00001"})\n'
        'MATCH (s:技能 {技能标准名: "Python"})\n'
        'MERGE (t)-[:拥有技能 {掌握程度: "熟练", 置信度: 0.95}]->(s);\n\n'
        '// 人才培养匹岗位（带匹配等级）\n'
        'MATCH (t:人才 {人才ID: "TALENT_00001"})\n'
        'MATCH (j:岗位 {岗位ID: "JOB_00001"})\n'
        'MERGE (t)-[:人岗匹配 {匹配等级: "高匹配", 匹配得分: 85}]->(j);'
    )

    doc.add_page_break()

    # ================================================================
    heading(doc, '五、核心查询语句（Cypher 中文版）')

    heading(doc, '5.1 查看某岗位的全部技能要求', 2)
    code_block(doc,
        'MATCH (j:岗位 {标准岗位名: "大模型应用开发工程师"})-[r:要求技能]->(s:技能)\n'
        'RETURN j.标准岗位名, s.技能标准名, r.权重\n'
        'ORDER BY r.权重 DESC'
    )

    heading(doc, '5.2 按技能分类统计岗位数量', 2)
    code_block(doc,
        'MATCH (j:岗位)-[:要求技能]->(s:技能)-[:属于分类]->(c:技能分类)\n'
        'RETURN c.分类名称, count(DISTINCT j) AS 岗位数量\n'
        'ORDER BY 岗位数量 DESC'
    )

    heading(doc, '5.3 查找与某人才最匹配的岗位（Top 10）', 2)
    code_block(doc,
        'MATCH (t:人才 {人才ID: "TALENT_00001"})-[:拥有技能]->(s:技能)<-[:要求技能]-(j:岗位)\n'
        'WITH j, count(s) AS 匹配技能数, collect(s.技能标准名) AS 匹配技能列表\n'
        'RETURN j.标准岗位名, j.企业名称, 匹配技能数, 匹配技能列表\n'
        'ORDER BY 匹配技能数 DESC LIMIT 10'
    )

    heading(doc, '5.4 计算人才与目标岗位的技能差距', 2)
    code_block(doc,
        '// 岗位要求但人才不具备的技能\n'
        'MATCH (j:岗位 {标准岗位名: "大模型应用开发工程师"})-[:要求技能]->(s:技能)\n'
        'WHERE NOT EXISTS {\n'
        '    MATCH (:人才 {人才ID: "TALENT_00001"})-[:拥有技能]->(s)\n'
        '}\n'
        'RETURN s.技能标准名 AS 缺失技能, s.所属分类 AS 分类'
    )

    heading(doc, '5.5 查找教授某缺失技能的课程', 2)
    code_block(doc,
        'MATCH (s:技能 {技能标准名: "RAG"})<-[:教授技能]-(c:课程)\n'
        'RETURN c.课程名称, c.提供平台, c.难度, c.时长'
    )

    heading(doc, '5.6 技能共现分析（发现新兴技能组合）', 2)
    code_block(doc,
        'MATCH (s1:技能)<-[:使用技术]-(p:技术项目)-[:使用技术]->(s2:技能)\n'
        'WHERE id(s1) < id(s2)\n'
        'RETURN s1.技能标准名 AS 技能A, s2.技能标准名 AS 技能B,\n'
        '       count(p) AS 共现次数\n'
        'ORDER BY 共现次数 DESC LIMIT 20'
    )

    heading(doc, '5.7 技能演化趋势查询', 2)
    code_block(doc,
        '// 按时间切片对比某岗位的技能变化\n'
        'MATCH (j:岗位 {标准岗位名: "Java开发工程师"})-[:要求技能]->(s:技能)\n'
        'RETURN s.技能标准名, j.时间切片\n'
        'ORDER BY j.时间切片'
    )

    heading(doc, '5.8 学习路径推荐（拓扑排序）', 2)
    code_block(doc,
        '// 从当前技能到目标技能的父技能链\n'
        'MATCH path = (current:技能 {技能标准名: "Python"})-[:父技能*1..3]->(target:技能 {技能标准名: "大语言模型"})\n'
        'RETURN [node in nodes(path) | node.技能标准名] AS 学习路径\n'
        'ORDER BY length(path) LIMIT 1'
    )

    doc.add_page_break()

    # ================================================================
    heading(doc, '六、Neo4j Browser 中文可视化配置')

    heading(doc, '6.1 节点颜色配置', 2)
    para(doc,
        '在 Neo4j Browser 中使用 Neo4j Browser 的 Style 功能（:style 命令），'
        '为不同中文标签的节点设置不同颜色，提升图谱可读性。'
    )
    code_block(doc,
        ':style\n'
        '{\n'
        '  "node": {\n'
        '    "岗位": {"color": "#E74C3C", "caption": "标准岗位名"},\n'
        '    "技能": {"color": "#3498DB", "caption": "技能标准名"},\n'
        '    "人才": {"color": "#2ECC71", "caption": "人才ID"},\n'
        '    "企业": {"color": "#F39C12", "caption": "企业名称"},\n'
        '    "课程": {"color": "#9B59B6", "caption": "课程名称"},\n'
        '    "证书": {"color": "#1ABC9C", "caption": "证书名称"},\n'
        '    "技术项目": {"color": "#E67E22", "caption": "项目名称"},\n'
        '    "论文": {"color": "#95A5A6", "caption": "论文标题"},\n'
        '    "技能分类": {"color": "#34495E", "caption": "分类名称"}\n'
        '  },\n'
        '  "relationship": {\n'
        '    "要求技能": {"color": "#E74C3C", "caption": "权重"},\n'
        '    "加分技能": {"color": "#E67E22", "caption": "权重"},\n'
        '    "拥有技能": {"color": "#2ECC71", "caption": "掌握程度"},\n'
        '    "人岗匹配": {"color": "#9B59B6", "caption": "匹配等级"}\n'
        '  }\n'
        '}'
    )

    heading(doc, '6.2 Neo4j Browser 展示效果说明', 2)
    para(doc,
        '应用上述配置后，在 Neo4j Browser 中执行任意查询，图谱将以以下方式呈现：\n'
        '  - 岗位节点：红色圆，标签显示标准岗位名（如 大模型应用开发工程师）\n'
        '  - 技能节点：蓝色圆，标签显示技能标准名（如 Python、大语言模型）\n'
        '  - 人才节点：绿色圆，标签显示人才ID\n'
        '  - 关系连线：箭头从源节点指向目标节点，不同关系类型用不同颜色区分\n'
        '  - 悬停节点时显示全部中文属性\n'
        '  - 点击关系连线时显示关系类型和属性（如 要求技能: 权重0.9）'
    )

    doc.add_page_break()

    # ================================================================
    heading(doc, '七、数据导入清单')

    heading(doc, '7.1 数据文件与节点/关系的对应', 2)
    table(doc,
        ['数据文件', '创建节点标签', '创建关系类型', '记录数'],
        [
            ['skill_ontology.json', '技能, 技能分类', '属于分类, 父技能', '545个技能, 10个分类'],
            ['jd_clean.csv', '岗位, 企业', '要求技能, 加分技能, 发布岗位', '2,500条'],
            ['gold_jd_set_reviewed.json', '岗位', '要求技能, 加分技能', '100条（人工复核）'],
            ['resume_clean.csv', '人才', '拥有技能', '180条'],
            ['gold_resume_set_reviewed.json', '人才', '拥有技能', '50条（人工复核）'],
            ['match_label_set_reviewed.json', '-', '人岗匹配', '100条（人工复核）'],
            ['course_data.csv', '课程', '教授技能', '384条'],
            ['certificate_data.csv', '证书', '认证技能', '101条'],
            ['github_trend.jsonl', '技术项目', '使用技术', '389条'],
            ['arxiv_trend.jsonl', '论文', '涉及技术', '359条'],
            ['negative_samples_reviewed.json', '-', '-', '50条（负样本测试）'],
        ],
    )

    heading(doc, '7.2 导入顺序建议', 2)
    for s in [
        '第1批：技能分类(10个) -> 技能(545个) -> 属于分类关系、父技能关系',
        '第2批：企业(~2000) -> 岗位(2500) -> 要求技能、加分技能、发布岗位关系',
        '第3批：人才(180) -> 拥有技能关系',
        '第4批：课程(384)、证书(101) -> 教授技能、认证技能关系',
        '第5批：技术项目(389)、论文(359) -> 使用技术、涉及技术关系',
        '第6批：人岗匹配关系(100条, 基于reviewed数据)',
        '每批导入后执行: MATCH (n) RETURN labels(n) LIMIT 1 验证导入成功',
    ]:
        bullet(doc, s)

    doc.add_page_break()

    # ================================================================
    heading(doc, '八、技术架构')
    table(doc,
        ['层次', '技术选型', '说明'],
        [
            ['图数据库', 'Neo4j Community 5.x', '原生图存储，Cypher 中文查询，中文可视化'],
            ['关系数据库', 'MySQL 8.0', '原始JD文本、用户数据、系统配置'],
            ['LLM服务', '讯飞星火大模型 API', '岗位定义生成、简历解析、自然语言解释'],
            ['后端', 'Node.js + Python FastAPI', 'Node负责API网关；Python负责图谱计算和AI'],
            ['前端', 'Vue3 + D3.js + Chart.js', '图谱可视化、管理后台、用户端'],
            ['部署', 'Docker Compose', '一键部署全部服务'],
        ],
    )

    heading(doc, '九、评估体系')
    table(doc,
        ['指标', '目标', '方法', '数据源'],
        [
            ['JD解析准确率', '>=90%', '技能抽取与gold_jd对比', 'gold_jd_set_reviewed.json'],
            ['简历提取准确率', '>=90%', '简历技能与gold_resume对比', 'gold_resume_set_reviewed.json'],
            ['人岗匹配准确率', '>=90%', '匹配等级与match_label对比', 'match_label_set_reviewed.json'],
            ['负样本识别率', '>=85%', '负样本判对率', 'negative_samples_reviewed.json'],
            ['技能映射覆盖率', '100%', '所有技能可映射到ontology', 'skill_ontology.json'],
        ],
    )
    para(doc,
        '测试方案含>=100条岗位JD及测试用例。人工评估基于300条reviewed数据，'
        '由5名审核人（余昭、徐赠贺、郭炫宇、邓佑杰、胡苗苗）完成。'
    )

    doc.add_page_break()

    heading(doc, '十、实施计划（4周）')
    table(doc,
        ['周次', '阶段', '任务', '产出'],
        [
            ['第1周', '环境+导入', 'Neo4j部署；Python批量导入节点和关系；Cypher查询验证', '可查询的中文图谱'],
            ['第2周', '新岗位发现', '技能共现分析；LLM生成岗位定义；人工审核界面', '新岗位候选列表'],
            ['第3周', '演化+匹配', '时间切片机制；变化检测；简历解析；人岗匹配API', '演化时间轴+匹配API'],
            ['第4周', '可视化+打磨', 'D3.js图谱可视化；学习路径；GraphRAG；文档', '完整可演示系统'],
        ],
    )

    doc.add_page_break()

    # ================================================================
    heading(doc, '附录A：完整Cypher建图脚本模板')
    para(doc, '以下是完整的 Neo4j 建图 Cypher 脚本骨架，可直接在 Neo4j Browser 或 cypher-shell 中执行：')
    code_block(doc,
        '// ===== 第1步：创建约束 =====\n'
        'CREATE CONSTRAINT 岗位ID约束 IF NOT EXISTS FOR (n:岗位) REQUIRE n.岗位ID IS UNIQUE;\n'
        'CREATE CONSTRAINT 技能名约束 IF NOT EXISTS FOR (n:技能) REQUIRE n.技能标准名 IS UNIQUE;\n'
        'CREATE CONSTRAINT 人才ID约束 IF NOT EXISTS FOR (n:人才) REQUIRE n.人才ID IS UNIQUE;\n'
        'CREATE CONSTRAINT 企业名约束 IF NOT EXISTS FOR (n:企业) REQUIRE n.企业名称 IS UNIQUE;\n'
        'CREATE CONSTRAINT 分类名约束 IF NOT EXISTS FOR (n:技能分类) REQUIRE n.分类名称 IS UNIQUE;\n\n'
        '// ===== 第2步：创建技能分类(10个) =====\n'
        'CREATE (:技能分类 {分类名称: "AI"});\n'
        'CREATE (:技能分类 {分类名称: "后端开发"});\n'
        'CREATE (:技能分类 {分类名称: "数据工程"});\n'
        'CREATE (:技能分类 {分类名称: "云计算"});\n'
        'CREATE (:技能分类 {分类名称: "前端开发"});\n'
        'CREATE (:技能分类 {分类名称: "AI智能体"});\n'
        'CREATE (:技能分类 {分类名称: "数据库"});\n'
        'CREATE (:技能分类 {分类名称: "物联网/嵌入式"});\n'
        'CREATE (:技能分类 {分类名称: "开发运维"});\n'
        'CREATE (:技能分类 {分类名称: "安全"});\n\n'
        '// ===== 第3步~第7步：使用 Python neo4j driver 批量导入 =====\n'
        '// 参考上文 4.3 节的 Python 代码模板\n\n'
        '// ===== 第8步：验证 =====\n'
        'MATCH (n) RETURN DISTINCT labels(n) AS 节点类型, count(n) AS 数量;\n'
        'MATCH ()-[r]->() RETURN DISTINCT type(r) AS 关系类型, count(r) AS 数量;'
    )

    doc.add_page_break()

    heading(doc, '附录B：节点和关系快速参考卡片')
    table(doc,
        ['类别', '中文名', '标识属性', '数量'],
        [
            ['节点', '岗位', '岗位ID', '2,500'],
            ['节点', '技能', '技能标准名', '545'],
            ['节点', '人才', '人才ID', '180'],
            ['节点', '企业', '企业名称', '~2,000'],
            ['节点', '课程', '课程ID', '384'],
            ['节点', '证书', '证书ID', '101'],
            ['节点', '技术项目', '项目ID', '389'],
            ['节点', '论文', '论文ID', '359'],
            ['节点', '技能分类', '分类名称', '10'],
            ['关系', '要求技能', '岗位->技能', '~15,000'],
            ['关系', '加分技能', '岗位->技能', '~8,000'],
            ['关系', '拥有技能', '人才->技能', '~1,200'],
            ['关系', '属于分类', '技能->技能分类', '545'],
            ['关系', '父技能', '技能->技能', '~200'],
            ['关系', '人岗匹配', '人才->岗位', '100'],
            ['关系', '教授技能', '课程->技能', '~1,500'],
            ['关系', '认证技能', '证书->技能', '~300'],
            ['关系', '发布岗位', '企业->岗位', '~2,500'],
            ['关系', '使用技术', '技术项目->技能', '~2,000'],
            ['关系', '涉及技术', '论文->技能', '~1,800'],
        ],
    )

    # 保存
    out = os.path.join(ROOT, '知识图谱详细设计文档_V2_中文图谱版.docx')
    doc.save(out)
    return out


if __name__ == '__main__':
    path = main()
    print(f'文档已生成: {path}')
